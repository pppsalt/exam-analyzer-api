"""
DOCX Output Generator
Creates landscape table with Unicode text + embedded diagram images in cells.
"""
import os
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

DIFFICULTY_COLORS = {
    "Easy": {"bg": "DCFCE7", "fg": RGBColor(0x16, 0x65, 0x34)},
    "Moderate": {"bg": "FEF9C3", "fg": RGBColor(0x85, 0x4D, 0x0E)},
    "Difficult": {"bg": "FEE2E2", "fg": RGBColor(0x99, 0x1B, 0x1B)},
}

HEADER_BG = "1E293B"
SUBTOPIC_NO_BG = "EFF6FF"

COL_WIDTHS_CM = [1.2, 1.2, 9.0, 3.5, 3.5, 5.0, 1.8, 6.5, 2.0]
HEADERS = [
    "S.No", "Q.No", "Question", "Chapter / Unit", "Topic",
    "Subtopic Name", "Sub.\nNo.", "Concept Tested", "Difficulty"
]


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    sd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex, qn("w:val"): "clear"
    })
    shading.append(sd)


def set_cell_text(cell, text: str, bold=False, italic=False, size=8.5,
                  color=None, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def add_image_to_cell(cell, image_path: str, max_width_cm=8.0):
    if not os.path.exists(image_path):
        return
    try:
        p = cell.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=Cm(max_width_cm))
    except Exception as e:
        logger.warning(f"Failed to embed image {image_path}: {e}")
        p = cell.add_paragraph()
        p.add_run(f"[Image: {os.path.basename(image_path)}]").italic = True


def generate(questions: list, metadata: dict, output_path: str) -> str:
    """
    Generate DOCX with landscape table.
    questions: list of dicts with all classification fields
    metadata: {exam_type, subject, model_used, paper_name, ...}
    """
    doc = Document()

    # Landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(metadata.get("paper_name", "Exam Paper Analysis"))
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_parts = []
    if metadata.get("exam_type"):
        info_parts.append(f"Exam: {metadata['exam_type']}")
    if metadata.get("subject"):
        info_parts.append(f"Subject: {metadata['subject']}")
    info_parts.append(f"Questions: {len(questions)}")
    if metadata.get("model_used"):
        model_name = metadata["model_used"].split("/")[-1]
        info_parts.append(f"Model: {model_name}")
    run = sub_p.add_run(" | ".join(info_parts))
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Table
    num_cols = len(HEADERS)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Set column widths
    for i, width in enumerate(COL_WIDTHS_CM):
        for cell in table.columns[i].cells:
            cell.width = Cm(width)

    # Header row
    hdr = table.rows[0]
    for i, header_text in enumerate(HEADERS):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header_text, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for idx, q in enumerate(questions):
        row = table.add_row()
        cells = row.cells

        sno = q.get("sno", idx + 1)
        qno = q.get("question_label", f"Q.{sno}")
        question_text = q.get("question_text", "")
        chapter = q.get("matched_unit_name") or q.get("topic", "")
        topic = q.get("topic", "")
        subtopic = q.get("matched_subtopic_name") or q.get("subtopic_name", "")
        sub_no = q.get("subtopic_number", "N/A")
        concept = q.get("concept_tested", "")
        difficulty = q.get("difficulty", "Moderate")

        # S.No
        set_cell_text(cells[0], str(sno), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Q.No
        set_cell_text(cells[1], qno, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Question — text + optional image
        set_cell_text(cells[2], question_text, size=8)
        diagram_paths = q.get("diagram_paths", [])
        if diagram_paths:
            for dp in diagram_paths:
                add_image_to_cell(cells[2], dp, max_width_cm=7.5)

        # Chapter
        set_cell_text(cells[3], chapter, size=8)
        # Topic
        set_cell_text(cells[4], topic, size=8)
        # Subtopic Name
        set_cell_text(cells[5], subtopic, bold=True, size=8)
        # Subtopic Number
        set_cell_text(cells[6], sub_no, bold=True, size=9,
                      color=RGBColor(0x1D, 0x4E, 0xD8),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cells[6], SUBTOPIC_NO_BG)
        # Concept Tested
        set_cell_text(cells[7], concept, italic=True, size=8)
        # Difficulty
        diff_cfg = DIFFICULTY_COLORS.get(difficulty, DIFFICULTY_COLORS["Moderate"])
        set_cell_text(cells[8], difficulty, bold=True, size=8,
                      color=diff_cfg["fg"], alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cells[8], diff_cfg["bg"])

    # Summary
    doc.add_paragraph()
    easy = sum(1 for q in questions if q.get("difficulty") == "Easy")
    mod = sum(1 for q in questions if q.get("difficulty") == "Moderate")
    diff = sum(1 for q in questions if q.get("difficulty") == "Difficult")

    summary_p = doc.add_paragraph()
    run = summary_p.add_run(f"Summary: ")
    run.font.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run = summary_p.add_run(
        f"Total: {len(questions)} | Easy: {easy} | Moderate: {mod} | Difficult: {diff}"
    )
    run.font.name = "Arial"
    run.font.size = Pt(9)

    # Footer info
    footer_p = doc.add_paragraph()
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    run = footer_p.add_run(f"Generated: {ts} | Model: {metadata.get('model_used', 'N/A')}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")
    return output_path
