import os
import re
import io
import json
import time
import uuid
import base64
import logging
import tempfile
import traceback
from datetime import datetime

import requests as http_requests
import pdfplumber
from PIL import Image
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from rapidfuzz import fuzz
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ─── Setup ───────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("exam-analyzer")

application = Flask(__name__)
CORS(application)

TEMP_DIR = tempfile.mkdtemp(prefix="exam_")
OUTPUT_DIR = os.path.join(TEMP_DIR, "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

# ─── Reference Data (loaded from JSON files if available) ────────────
REF_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reference_data")
_ref_cache = {}

def load_reference(exam_type, subject):
    key = f"{exam_type}_{subject}"
    if key in _ref_cache:
        return _ref_cache[key]
    filepath = os.path.join(REF_DIR, f"{key}.json")
    if not os.path.exists(filepath):
        _ref_cache[key] = []
        return []
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            rows = json.load(f)
        _ref_cache[key] = rows
        logger.info(f"Loaded {len(rows)} subtopics for {key}")
        return rows
    except Exception:
        _ref_cache[key] = []
        return []

def match_subtopic(ai_subtopic, ai_topic, exam_type, subject):
    refs = load_reference(exam_type, subject)
    if not refs:
        return {"subtopic_number": "N/A", "matched_name": None, "confidence": 0, "unit_name": None}
    best_match, best_score = None, 0
    for r in refs:
        for text in [r["subtopic_name"], f"{r['unit_name']}: {r['subtopic_name']}"]:
            for query in [ai_subtopic, f"{ai_topic}: {ai_subtopic}", f"{ai_topic} {ai_subtopic}"]:
                score = fuzz.token_sort_ratio(query.lower(), text.lower())
                if score > best_score:
                    best_score = score
                    best_match = r
    if best_match and best_score >= 60:
        return {"subtopic_number": best_match["subtopic_number"], "matched_name": best_match["subtopic_name"],
                "unit_name": best_match["unit_name"], "confidence": best_score}
    return {"subtopic_number": "N/A", "matched_name": None, "confidence": best_score, "unit_name": None}

# ─── PDF Extraction ──────────────────────────────────────────────────
def extract_pdf(pdf_path, output_dir):
    text = ""
    images = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text(x_tolerance=2, y_tolerance=2)
                if t:
                    text += t + "\n"
                for j, img in enumerate(page.images[:5]):
                    try:
                        bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                        cropped = page.within_bbox(bbox).to_image(resolution=150)
                        img_path = os.path.join(output_dir, f"img_p{i+1}_{j+1}.png")
                        cropped.save(img_path)
                        images.append({"page": i+1, "path": img_path})
                    except Exception:
                        pass
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")

    # Parse questions
    questions = []
    lines = text.split("\n")
    current_q = None
    q_pattern = re.compile(r'^(?:Q\.?\s*)?(\d{1,3})[\.\)\:]?\s+(.+)', re.IGNORECASE)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        m = q_pattern.match(line)
        if m:
            if current_q:
                questions.append(current_q)
            num = int(m.group(1))
            current_q = {"number": num, "text": m.group(2), "label": f"Q.{num}", "diagram_paths": []}
        elif current_q:
            current_q["text"] += " " + line

    if current_q:
        questions.append(current_q)

    # Auto-detect exam type and subject
    text_lower = text.lower()
    exam_type = "UNKNOWN"
    if "jee" in text_lower:
        exam_type = "JEE"
    elif "neet" in text_lower:
        exam_type = "NEET"

    subject = "UNKNOWN"
    for s in ["Physics", "Chemistry", "Mathematics", "Biology"]:
        if s.lower() in text_lower:
            subject = s
            break

    return {"questions": questions, "exam_type": exam_type, "subject": subject, "full_text": text}

# ─── AI Analysis ─────────────────────────────────────────────────────
def ai_analyze(questions, model_id, exam_type, subject):
    api_key = os.environ.get("OPENROUTER_API_KEY", "")
    if not api_key:
        raise ValueError("OPENROUTER_API_KEY not set")

    q_text = ""
    for q in questions[:50]:
        q_text += f"Q{q['number']}: {q['text'][:500]}\n"

    prompt = f"""You are an expert {exam_type} exam analyst for {subject}.
Analyze each question below and return a JSON array. For each question return:
- sno (question number)
- question_text (first 200 chars)
- subject ("{subject}")
- topic (main chapter/topic)
- subtopic_name (specific subtopic)
- concept_tested (what concept is being tested)
- difficulty ("Easy", "Moderate", or "Hard")
- has_diagram (true/false - guess from question text)
- diagram_description (if diagram likely, describe what it might show)

Questions:
{q_text}

Return ONLY a valid JSON array, nothing else. No markdown, no explanation."""

    headers = {
        "Authorization": "Bearer " + api_key,
        "Content-Type": "application/json",
        "HTTP-Referer": "https://exam-analyzer.onrender.com",
    }

    payload = {
        "model": model_id,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 4000,
    }

    for attempt in range(3):
        try:
            resp = http_requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=90)
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]

            # Clean markdown fences
            content = content.strip()
            if content.startswith("```"):
                content = re.sub(r'^```(?:json)?\s*', '', content)
                content = re.sub(r'\s*```$', '', content)

            result = json.loads(content)
            if isinstance(result, list):
                return {"questions": result}
            elif isinstance(result, dict) and "questions" in result:
                return result
            else:
                return {"questions": [result]}

        except Exception as e:
            logger.error(f"AI attempt {attempt+1} failed: {e}")
            if attempt == 2:
                raise ValueError(f"AI analysis failed after 3 attempts: {e}")
            time.sleep(2)

# ─── DOCX Generation ─────────────────────────────────────────────────
def generate_docx(questions, metadata, output_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(9)
    style.font.name = "Calibri"

    doc.add_heading(f"{metadata.get('paper_name', 'Exam')} — Analysis Report", level=1)
    doc.add_paragraph(f"Exam: {metadata.get('exam_type', 'N/A')} | Subject: {metadata.get('subject', 'N/A')} | "
                      f"Model: {metadata.get('model_used', 'N/A')} | Questions: {len(questions)}")

    headers = ["#", "Question", "Topic", "Subtopic", "Sub#", "Concept", "Diff"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)

    for q in questions:
        row = table.add_row()
        vals = [
            str(q.get("sno", "")),
            (q.get("question_text", "") or "")[:120],
            q.get("topic", ""),
            q.get("subtopic_name", ""),
            q.get("subtopic_number", "N/A"),
            q.get("concept_tested", ""),
            q.get("difficulty", ""),
        ]
        for i, v in enumerate(vals):
            row.cells[i].text = str(v)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    doc.save(output_path)

# ─── XLSX Generation ─────────────────────────────────────────────────
def generate_xlsx(questions, metadata, output_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Analysis"

    headers = ["#", "Question", "Subject", "Topic", "Subtopic", "Sub#", "Concept", "Difficulty", "Diagram"]
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=10)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font

    for row_idx, q in enumerate(questions, 2):
        ws.cell(row=row_idx, column=1, value=q.get("sno", ""))
        ws.cell(row=row_idx, column=2, value=(q.get("question_text", "") or "")[:200])
        ws.cell(row=row_idx, column=3, value=q.get("subject", ""))
        ws.cell(row=row_idx, column=4, value=q.get("topic", ""))
        ws.cell(row=row_idx, column=5, value=q.get("subtopic_name", ""))
        ws.cell(row=row_idx, column=6, value=q.get("subtopic_number", "N/A"))
        ws.cell(row=row_idx, column=7, value=q.get("concept_tested", ""))
        ws.cell(row=row_idx, column=8, value=q.get("difficulty", ""))
        ws.cell(row=row_idx, column=9, value="Yes" if q.get("has_diagram") else "No")

    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 25
    ws.column_dimensions["G"].width = 30
    wb.save(output_path)

# ─── Routes ──────────────────────────────────────────────────────────
@application.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "timestamp": datetime.now().isoformat()})

@application.route("/models", methods=["GET"])
def get_models():
    models = [
        {"model_id": "google/gemini-2.5-flash", "display_name": "Gemini 2.5 Flash", "provider": "Google", "supports_vision": True},
        {"model_id": "google/gemini-2.5-pro", "display_name": "Gemini 2.5 Pro", "provider": "Google", "supports_vision": True},
        {"model_id": "anthropic/claude-sonnet-4", "display_name": "Claude Sonnet 4", "provider": "Anthropic", "supports_vision": True},
        {"model_id": "openai/gpt-4o", "display_name": "GPT-4o", "provider": "OpenAI", "supports_vision": True},
        {"model_id": "openai/o3-mini", "display_name": "o3-mini", "provider": "OpenAI", "supports_vision": False},
        {"model_id": "deepseek/deepseek-r1", "display_name": "DeepSeek R1", "provider": "DeepSeek", "supports_vision": False},
    ]
    return jsonify({"models": models})

@application.route("/analyze", methods=["POST"])
def analyze():
    pdf_file = request.files.get("pdf_file")
    model_id = request.form.get("model_id", "google/gemini-2.5-flash")
    forced_exam = request.form.get("exam_type", "")
    forced_subject = request.form.get("subject", "")
    upload_id = request.form.get("upload_id", "")

    if not pdf_file:
        return jsonify({"status": "failed", "error": "No PDF file uploaded"}), 400

    start_time = time.time()
    job_id = str(uuid.uuid4())[:8]
    job_dir = os.path.join(TEMP_DIR, job_id)
    os.makedirs(job_dir, exist_ok=True)

    pdf_path = os.path.join(job_dir, pdf_file.filename)
    pdf_file.save(pdf_path)

    try:
        # Step 1: Extract
        logger.info(f"[{job_id}] Extracting PDF")
        extraction = extract_pdf(pdf_path, job_dir)
        questions = extraction["questions"]
        exam_type = forced_exam or extraction["exam_type"]
        subject = forced_subject or extraction["subject"]

        if not questions:
            raise ValueError("No questions detected in PDF")

        logger.info(f"[{job_id}] Found {len(questions)} questions, exam={exam_type}, subject={subject}")

        # Step 2: AI Analysis
        logger.info(f"[{job_id}] AI Analysis with {model_id}")
        ai_result = ai_analyze(questions, model_id, exam_type, subject)
        ai_questions = ai_result["questions"]

        # Step 3: Subtopic matching
        for q in ai_questions:
            subj = q.get("subject", subject)
            result = match_subtopic(q.get("subtopic_name", ""), q.get("topic", ""), exam_type, subj)
            q["subtopic_number"] = result["subtopic_number"]
            q["match_confidence"] = result["confidence"]
            q["matched_subtopic_name"] = result["matched_name"]
            q["matched_unit_name"] = result["unit_name"]

        # Step 4: Generate outputs
        paper_name = os.path.splitext(pdf_file.filename)[0]
        metadata = {"paper_name": paper_name, "exam_type": exam_type, "subject": subject, "model_used": model_id}

        docx_filename = f"{paper_name}_{job_id}.docx"
        xlsx_filename = f"{paper_name}_{job_id}.xlsx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        xlsx_path = os.path.join(OUTPUT_DIR, xlsx_filename)

        generate_docx(ai_questions, metadata, docx_path)
        generate_xlsx(ai_questions, metadata, xlsx_path)

        elapsed = time.time() - start_time
        logger.info(f"[{job_id}] Done in {elapsed:.1f}s")

        return jsonify({
            "status": "completed", "job_id": job_id, "upload_id": upload_id,
            "questions_count": len(ai_questions), "exam_type": exam_type, "subject": subject,
            "model_used": model_id, "processing_time": round(elapsed, 1),
            "docx_filename": docx_filename, "xlsx_filename": xlsx_filename,
            "docx_url": f"/download/{docx_filename}", "xlsx_url": f"/download/{xlsx_filename}",
            "questions": ai_questions,
        })

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"[{job_id}] Failed: {e}\n{traceback.format_exc()}")
        return jsonify({"status": "failed", "error": str(e), "processing_time": round(elapsed, 1)}), 500

@application.route("/download/<filename>", methods=["GET"])
def download(filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404
    return send_file(filepath, as_attachment=True)

# ─── Entry Point ─────────────────────────────────────────────────────
app = application

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5050))
    application.run(host="0.0.0.0", port=port, debug=False)
